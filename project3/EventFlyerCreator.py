#!/usr/bin/python3

from docx import Document 

##Create the flyer document

document = Document()



##Greeting the user to the app

print('Welcome to the Event Flyer Generator v1')

print()
print()

##Enter title for event

Event_Title = input('Please choose a title for this Event: ')

print()


##Enter event Description

Event_Description = input('Please input a brief description for this event: ')

print()


##Enter the host for the event

Event_Host = input('Please enter the Host for the Event: ' )

print()


##This will create a reusable function to allow the user to make selections

def get_choice(available_choices, choice_type):
    while True:
        print()
        print()
        print(f'Please choose a {choice_type}')
    
    
        for choice in available_choices:
            print(f'[*] ' + choice)
        

        print()

        user_choice = input( 'Choice: ')
        if user_choice.lower() in available_choices:
            return user_choice
        else:
            print()
            print('Please make a valid Selection')





##Time selection for Event

Available_Event_Times = ( '1', '2', '3', '4', '5')

##Allows the user to select A time

Selected_Event_Time = get_choice(Available_Event_Times, 'Time (PM)')


##Event locations

Available_Event_Locations = ('123 college', 'studio 123' , 'york park')


##Allows the user to select a location

Selected_Event_Location = get_choice(Available_Event_Locations, 'Location')


##Event Price

Available_Event_Prices = ('$500' , '$200' , '$100' , '$50')

Selected_Event_Price = get_choice(Available_Event_Prices, 'Price')


#Save the image as a variable and allow the user to select the variable


Available_Event_Images = ('sports', 'books', 'cars', 'theater')



Selected_Event_Image = get_choice(Available_Event_Images, 'Event Image')

if Selected_Event_Image == 'sports':
    Event_image = "athlete.jpg"
elif Selected_Event_Image == 'books':
    Event_image = "books.jpg"
elif Selected_Event_Image == 'cars':
    Event_image = "carshow.jpg"
else:
    Event_image = "concert.jpg"



##Adding content to word document

document.add_heading(Event_Title, 0)

Description = document.add_paragraph("Description: " + Event_Description)

Times = document.add_paragraph("Time: ")
Times.add_run( Selected_Event_Time + " P.M").bold = True

Location = document.add_paragraph("Location: ")
Location.add_run( Selected_Event_Location)


Price = document.add_paragraph ("Price: ")
Price.add_run(Selected_Event_Price+ "+ $50 per additional guest").bold = True

Image = document.add_picture(Event_image)

document.save("EventFlyer.docx")


print("Your flyer has been created, Please navigate to the folder containing the script")
